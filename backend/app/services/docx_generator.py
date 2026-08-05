from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from app.models.enums import CVType
from app.models.responses import (ContactBlock, ResumeSection, TailorResponse, )


def apply_cv_type_styling(doc: Document, cv_type: CVType) -> None:
    section = doc.sections[0]

    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = doc.styles["Normal"].font
    heading = doc.styles["Heading 1"].font

    if cv_type == CVType.TECH:
        normal.name = "Arial"
        normal.size = Pt(11)

        heading.name = "Arial"
        heading.color.rgb = RGBColor(0x00, 0x33, 0x66)
    else:
        normal.name = "Garamond"
        normal.size = Pt(11)

        heading.name = "Garamond"
        heading.color.rgb = RGBColor(0x33, 0x33, 0x33)

    heading.bold = True
    heading.size = Pt(14)


def render_contact_block(doc: Document, contact_info: ContactBlock) -> None:
    paragraph = doc.add_paragraph()

    run = paragraph.add_run(contact_info.name)
    run.bold = True
    run.font.size = Pt(18)

    details = []

    if contact_info.email:
        details.append(contact_info.email)

    if contact_info.phone:
        details.append(contact_info.phone)

    if contact_info.location:
        details.append(contact_info.location)

    if details:
        doc.add_paragraph(" | ".join(details))

    for link in contact_info.links:
        doc.add_paragraph(link)


def render_section(doc: Document, section: ResumeSection) -> None:
    doc.add_heading(section.section_name, level=1)

    for entry in section.entries:
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(entry.title)
        title_run.bold = True
        title_run.font.size = Pt(12)

        if entry.subtitle:
            subtitle_paragraph = doc.add_paragraph()
            subtitle_run = subtitle_paragraph.add_run(entry.subtitle)
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(11)

        for bullet in entry.bullets:
            bullet_paragraph = doc.add_paragraph(
                bullet,
                style="List Bullet",
            )
            bullet_paragraph.paragraph_format.space_after = Pt(2)


def generate_docx(tailor_response: TailorResponse) -> bytes:
    doc = Document()

    apply_cv_type_styling(doc, tailor_response.cv_type)

    render_contact_block(doc, tailor_response.contact_info)

    for section in sorted(
        tailor_response.sections,
        key=lambda section: section.order_index,
    ):
        render_section(doc, section)

    buffer = BytesIO()

    doc.save(buffer)

    return buffer.getvalue()
