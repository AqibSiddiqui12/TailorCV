import io

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

from gemini import TailoredCV

def add_bottom_border(paragraph):
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4') 
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pbdr.append(bottom)
    pPr.append(pbdr)

def add_flex_row(doc, left_text: str, right_text: str, bold_left=True):
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    r = p.add_run(left_text)
    r.bold = bold_left
    p.add_run(f"\t{right_text}")
    p.paragraph_format.space_after = Pt(2)
    return p

def create_docx(cv_data: TailoredCV) -> io.BytesIO:
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(cv_data.name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_p.paragraph_format.space_after = Pt(2)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.paragraph_format.space_after = Pt(12)
    contact_p.paragraph_format.line_spacing = 1.1
    
    contact_parts = []
    if cv_data.contact.location: contact_parts.append(cv_data.contact.location)
    if cv_data.contact.phone: contact_parts.append(cv_data.contact.phone)
    if cv_data.contact.email: contact_parts.append(cv_data.contact.email)
        
    if contact_parts:
        contact_p.add_run(" • ".join(contact_parts) + "\n")
    if cv_data.contact.links:
        contact_p.add_run(" • ".join(cv_data.contact.links))

    def add_section_header(title: str):
        h = doc.add_paragraph()
        r = h.add_run(title.upper())
        r.bold = True
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True # Fixes the page break issue!
        add_bottom_border(h)

    if cv_data.summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph(cv_data.summary)
        p.paragraph_format.space_after = Pt(6)

    if cv_data.experience:
        add_section_header("Professional Experience")
        for exp in cv_data.experience:
            add_flex_row(doc, exp.company.upper(), exp.dates, bold_left=True)
            title_p = doc.add_paragraph()
            title_p.add_run(exp.title).italic = True
            title_p.paragraph_format.space_after = Pt(2)
            title_p.paragraph_format.keep_with_next = True
            for bullet in exp.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if cv_data.projects:
        add_section_header("Projects")
        for proj in cv_data.projects:
            add_flex_row(doc, proj.name.upper(), "", bold_left=True)
            for bullet in proj.bullets:
                bp = doc.add_paragraph(bullet, style="List Bullet")
                bp.paragraph_format.left_indent = Inches(0.25)
                bp.paragraph_format.space_after = Pt(2)
            doc.add_paragraph().paragraph_format.space_after = Pt(4)

    if cv_data.education:
        add_section_header("Education")
        for edu in cv_data.education:
            add_flex_row(doc, edu.institution.upper(), edu.dates, bold_left=True)
            deg_p = doc.add_paragraph(edu.degree)
            deg_p.paragraph_format.space_after = Pt(6)

    if cv_data.skills:
        add_section_header("Technical Skills & Additional")
        for skill_cat in cv_data.skills:
            p = doc.add_paragraph()
            p.add_run(f"{skill_cat.category}: ").bold = True
            p.add_run(", ".join(skill_cat.items))
            p.paragraph_format.space_after = Pt(2)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream